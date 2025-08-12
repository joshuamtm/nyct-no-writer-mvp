"""
Document processing service for extracting text from PDFs and Word documents.
"""
import io
from typing import Optional
import pypdf
from docx import Document
import pdfplumber
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles text extraction from various document formats."""
    
    @staticmethod
    async def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF using multiple methods for better accuracy.
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            Extracted text as string
        """
        text = ""
        
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # If pdfplumber didn't get much, try pypdf
            if len(text.strip()) < 100:
                pdf_reader = pypdf.PdfReader(io.BytesIO(file_content))
                pypdf_text = ""
                for page in pdf_reader.pages:
                    pypdf_text += page.extract_text() + "\n"
                
                if len(pypdf_text.strip()) > len(text.strip()):
                    text = pypdf_text
                    
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            # Fallback to basic pypdf extraction
            try:
                pdf_reader = pypdf.PdfReader(io.BytesIO(file_content))
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            except Exception as fallback_error:
                logger.error(f"Fallback PDF extraction failed: {fallback_error}")
                raise
        
        return text.strip()
    
    @staticmethod
    async def extract_text_from_docx(file_content: bytes) -> str:
        """
        Extract text from Word document.
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Extracted text as string
        """
        try:
            doc = Document(io.BytesIO(file_content))
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return "\n".join(text)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
    
    @staticmethod
    async def extract_text(file_content: bytes, content_type: str) -> str:
        """
        Extract text from document based on content type.
        
        Args:
            file_content: File content as bytes
            content_type: MIME type of the file
            
        Returns:
            Extracted text as string
        """
        if content_type == "application/pdf":
            return await DocumentProcessor.extract_text_from_pdf(file_content)
        elif content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]:
            return await DocumentProcessor.extract_text_from_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {content_type}")
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                # Remove multiple spaces
                line = ' '.join(line.split())
                cleaned_lines.append(line)
        
        # Join with single newlines
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove very long sequences of special characters
        import re
        cleaned_text = re.sub(r'[_\-=]{10,}', '', cleaned_text)
        
        return cleaned_text